# Streamlit app: Excel (ID + HTML) -> Word (.docx)
# This reverses your Word->HTML->Excel pipeline as closely as possible

import streamlit as st
from docx import Document
from openpyxl import load_workbook
from io import BytesIO
from datetime import datetime
from bs4 import BeautifulSoup

# ============================
# HTML -> Word helpers
# ============================

def add_html_to_doc(doc, html):
    """Convert a limited HTML subset back into Word paragraphs.
    If HTML contains no tags, treat it as plain text.
    """

    # If the cell is plain text (no HTML tags), just write it
    if "<" not in str(html) or ">" not in str(html):
        p = doc.add_paragraph()
        p.add_run(str(html))
        return

    soup = BeautifulSoup(html, "html.parser")

    for element in soup.contents:
        if element.name == "p":
            p = doc.add_paragraph()
            add_inline_runs(p, element)

        elif element.name == "ul":
            for li in element.find_all("li", recursive=False):
                p = doc.add_paragraph(style="List Bullet")
                add_inline_runs(p, li)


def add_inline_runs(paragraph, element):
    """Safely convert inline HTML to Word runs without duplication."""
    for child in element.children:
        # Bold / strong text
        if getattr(child, "name", None) in ("b", "strong"):
            run = paragraph.add_run(child.get_text())
            run.bold = True

        # Plain text node
        elif child.name is None:
            text = child.strip()
            if text:
                paragraph.add_run(text)

        # Any other nested tag (future-proofing)
        else:
            add_inline_runs(paragraph, child)



# ============================
# Excel -> Word
# ============================

def excel_to_word(excel_file):
    wb = load_workbook(excel_file)
    ws = wb.active

    doc = Document()

    # Skip header row
    for row in ws.iter_rows(min_row=2, values_only=True):
        product_id, html = row
        if not product_id:
            continue

        # ID as heading / separator
        id_para = doc.add_paragraph(str(product_id))
        id_para.runs[0].bold = True

        # Required 20000 separator
        doc.add_paragraph("20000")

        if html:
            add_html_to_doc(doc, html)

        doc.add_page_break()

    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output


# ============================
# Streamlit UI
# ============================

st.title("Excel to Word Converter")
st.write("Upload an Excel file generated from the Word→HTML tool")

uploaded_file = st.file_uploader("Choose an Excel (.xlsx) file", type=["xlsx"])

if uploaded_file:
    st.success(f"File uploaded: {uploaded_file.name}")

    if st.button("Convert to Word"):
        with st.spinner("Converting..."):
            word_file = excel_to_word(uploaded_file)

        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_filename = f"reconstructed_{timestamp}.docx"

        st.download_button(
            label="Download Word Document",
            data=word_file,
            file_name=output_filename,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

        st.success("Conversion complete!")
